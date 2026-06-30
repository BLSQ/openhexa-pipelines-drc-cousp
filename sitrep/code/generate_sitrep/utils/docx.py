from __future__ import annotations

import contextlib
from collections.abc import Callable

from docx.document import Document as DocumentT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.xmlchemy import BaseOxmlElement
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

FONT = "Arial Narrow"
CENTER = WD_ALIGN_PARAGRAPH.CENTER
# Fond des en-têtes et lignes « Total » des tableaux de données (gris).
HEADER_FILL = "BFBFBF"
# Bordures « filets horizontaux » des tableaux de données : contour haut/bas
BORDER_THICK = 12
BORDER_THIN = 2
_TBLPR_AFTER_BORDERS = ("w:shd", "w:tblLayout", "w:tblCellMar", "w:tblLook")


def norm(s: str) -> str:
    """Normalise un libellé : apostrophe droite, espaces/retours compactés.

    Returns:
        La chaîne normalisée en minuscules, avec les espaces et retours compactés.
    """
    return " ".join(s.replace("’", "'").split()).lower()  # noqa: RUF001


def set_cell(
    cell: _Cell,
    text: str,
    *,
    bold: bool = False,
    align: WD_ALIGN_PARAGRAPH | None = None,
    color: str | None = None,
    size: int = 10,
) -> None:
    """Set text and formatting for a table cell.

    Args:
        cell: The docx table cell to update.
        text: The text content to insert into the cell.
        bold: Whether the text should be bold.
        align: Optional paragraph alignment for the cell.
        color: Optional RGB color string for the text.
        size: Font size in points (10 par défaut).
    """
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = FONT
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell: _Cell, fill: str = HEADER_FILL) -> None:
    """Applique une couleur de fond (hex sans ``#``) à une cellule de tableau.

    Args:
        cell: La cellule à colorer.
        fill: La couleur de remplissage hexadécimale (``BFBFBF`` par défaut).
    """
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def _border_el(side: str, sz: int) -> BaseOxmlElement:
    """Construit un élément de bordure ``w:<side>`` (``none`` si ``sz <= 0``).

    Returns:
        L'élément XML de bordure prêt à insérer.
    """
    el = OxmlElement(f"w:{side}")
    if sz <= 0:
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
    else:
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
    el.set(qn("w:space"), "0")
    el.set(qn("w:color"), "auto")
    return el


def set_table_borders(
    table: Table,
    *,
    thick: int = BORDER_THICK,
    thin: int = BORDER_THIN,
) -> None:
    """Applique le style « filets horizontaux » aux bordures d'un tableau.

    Args:
        table: Le tableau à border.
        thick: Épaisseur (douzièmes de pt) du contour haut/bas.
        thin: Épaisseur (douzièmes de pt) de la grille interne (0 = aucune).
    """
    tbl_pr = table._tbl.tblPr
    for old in tbl_pr.findall(qn("w:tblBorders")):
        tbl_pr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for side, sz in (
        ("top", thick),
        ("left", 0),
        ("bottom", thick),
        ("right", 0),
        ("insideH", thin),
        ("insideV", thin),
    ):
        borders.append(_border_el(side, sz))

    for tag in _TBLPR_AFTER_BORDERS:
        ref = tbl_pr.find(qn(tag))
        if ref is not None:
            ref.addprevious(borders)
            break
    else:
        tbl_pr.append(borders)


def set_cell_border_bottom(cell: _Cell, sz: int = BORDER_THICK) -> None:
    """Pose un filet épais sous une cellule (séparateur sous l'en-tête).

    Args:
        cell: La cellule à souligner.
        sz: Épaisseur du filet (douzièmes de pt).
    """
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is not None:
        borders.remove(bottom)
    borders.append(_border_el("bottom", sz))


class Cursor:
    """Insère des éléments XML à la suite d'un élément d'ancrage."""

    def __init__(self, anchor_element: BaseOxmlElement) -> None:
        self._el = anchor_element

    def add(self, new_element: BaseOxmlElement) -> None:
        """Insert a new XML element after the current anchor and move the cursor.

        Args:
            new_element: The XML element to insert after the current cursor position.
        """
        self._el.addnext(new_element)
        self._el = new_element


def para(
    doc: DocumentT,
    text: str = "",
    *,
    bold: bool = False,
    size: int = 10,
    italic: bool = False,
    color: str | None = None,
    align: WD_ALIGN_PARAGRAPH | None = None,
) -> Paragraph:
    """Add a paragraph to the document with optional formatting.

    Args:
        doc: Document to add the paragraph to.
        text: Text content for the paragraph.
        bold: Whether the text should be bold.
        size: Font size in points.
        italic: Whether the text should be italic.
        color: Optional RGB color string for the text.
        align: Optional paragraph alignment.

    Returns:
        The created paragraph.
    """
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = FONT
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
    return p


def bullet(doc: DocumentT, text: str) -> Paragraph:
    """Puce, en repli manuel si le style « List Bullet » est absent.

    Returns:
        The created paragraph.
    """
    try:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(text)
    except KeyError:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        run = p.add_run(f"•  {text}")
    run.font.name = FONT
    run.font.size = Pt(10)
    return p


def table(doc: DocumentT, headers: list[str], rows: list[list]) -> Table:
    """Create a table in the document with the given headers and rows.

    Args:
        doc: The document where the table will be inserted.
        headers: A list of header titles for the table columns.
        rows: A list of rows, each row being a list of cell values.

    Returns:
        The created table object.
    """
    t = doc.add_table(rows=1, cols=len(headers))
    with contextlib.suppress(KeyError):
        t.style = "Table Grid"

    # Bordures « filets horizontaux » (contour haut/bas, grille interne fine).
    set_table_borders(t)

    # En-tête : texte gras sur fond gris + filet épais en dessous.
    for j, h in enumerate(headers):
        set_cell(t.cell(0, j), h, bold=True, align=CENTER)
        set_cell_shading(t.cell(0, j))
        set_cell_border_bottom(t.cell(0, j))

    for r in rows:
        cells = t.add_row().cells
        is_total = str(r[0]).strip().lower() == "total"
        for j, cell in enumerate(cells):
            set_cell(cell, r[j], bold=is_total, align=CENTER if j else None)

            # Ligne de total : même fond gris que l'en-tête.
            if is_total:
                set_cell_shading(cell)

    # Filet de clôture : bottom explicite sur la dernière ligne — Word ne rend
    if len(t.rows) > 1:
        for cell in t.rows[-1].cells:
            set_cell_border_bottom(cell)
    return t


def marker_paragraph(doc: DocumentT, token: str) -> Paragraph | None:
    """Retourne le premier paragraphe dont le texte commence par ``token``.

    Returns:
        Paragraph | None: Le paragraphe marqueur, ou ``None`` s'il est absent.
    """
    for p in doc.paragraphs:
        if p.text.strip().startswith(token):
            return p
    return None


def replace_marker(doc: DocumentT, token: str, fill: Callable[[Cursor], None]) -> None:
    """Insère le contenu produit par ``fill(cursor)`` puis supprime le marqueur."""
    marker = marker_paragraph(doc, token)
    if marker is None:
        return
    cur = Cursor(marker._p)
    fill(cur)
    marker._p.getparent().remove(marker._p)


def find_table(doc: DocumentT, predicate: Callable[[Table], bool]) -> Table | None:
    """Retourne la première table satisfaisant ``predicate``.

    Returns:
        Table | None: La table trouvée, ou ``None`` si aucune ne correspond.
    """
    for t in doc.tables:
        try:
            if predicate(t):
                return t
        except Exception:
            continue
    return None
