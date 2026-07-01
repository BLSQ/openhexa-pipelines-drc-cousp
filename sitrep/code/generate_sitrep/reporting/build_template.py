from __future__ import annotations

from pathlib import Path

import config
from docx import Document
from docx.document import Document as DocumentT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.table import _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run

_FONT = "Arial Narrow"
_RED = config.ACCENT_RED

# Libellés des chiffres-clés (ordre d'affichage du tableau de situation globale).
KPI_LABELS = [
    "Cumul cas confirmés",
    "Cumul décès parmi les confirmés",
    "Cas suspects en cours d’investigation",  # noqa: RUF001
    "Cas suspects en isolement",
    "Cas confirmés actifs",
    "Guéris",
    "Taux de suivi de contacts",
]


# En-têtes du tableau de synthèse par province (alignés sur le SitRep officiel).
PROVINCE_HEADERS = [
    "Provinces touchées",
    "Cas confirmés",
    "Décès confirmés",
    "Cas suspects en cours d’investigation",  # noqa: RUF001
    "Cas suspects en isolement",
    "Cas confirmés actifs",
    "Guéris du jour",
]

IDENTITY_ROWS = [
    "Provinces touchées",
    "Zones de Santé touchées",
    "Date de rapportage",
    "Date de publication",
]

# Sous-sections des actions de réponse (titre affiché, marqueur).
ACTION_SUBSECTIONS = [
    ("Coordination", "[[ACTIONS_COORDINATION]]"),
    ("Surveillance épidémiologique", "[[ACTIONS_SURVEILLANCE]]"),
    ("Laboratoire", "[[ACTIONS_LABORATOIRE]]"),
    (
        "Communication sur les risques et engagement communautaire (CREC)",
        "[[ACTIONS_CREC]]",
    ),
    ("Prévention et contrôle de l'infection (PCI / WASH)", "[[ACTIONS_PCI_WASH]]"),
    ("Logistique", "[[ACTIONS_LOGISTIQUE]]"),
    ("Sécurité", "[[ACTIONS_SECURITE]]"),
]


def _set_run(
    run: Run,
    *,
    bold: bool = False,
    size: int = 11,
    color: str | None = None,
    italic: bool = False,
) -> None:
    run.bold = bold
    run.italic = italic
    run.font.name = _FONT
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _title_line(
    doc: DocumentT,
    text: str,
    *,
    size: int,
    bold: bool = True,
    color: str = _RED,
) -> Paragraph:
    """Ajoute une ligne de titre centrée.

    Returns:
        Paragraph: Le paragraphe de titre créé.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    _set_run(p.add_run(text), bold=bold, size=size, color=color)
    return p


def _heading(doc: DocumentT, text: str) -> Paragraph:
    """Titre de section rouge (style Heading 1 du gabarit par défaut).

    Returns:
        Paragraph: Le paragraphe de titre créé.
    """
    p = doc.add_paragraph(style="Heading 1")
    run = p.add_run(text)
    _set_run(run, bold=True, size=14, color=_RED)
    return p


def _subheading(doc: DocumentT, text: str) -> Paragraph:
    """Sous-titre de section (style Heading 2).

    Returns:
        Paragraph: Le paragraphe de sous-titre créé.
    """
    p = doc.add_paragraph(style="Heading 2")
    run = p.add_run(text)
    _set_run(run, bold=True, size=12, color="7B241C")
    return p


def _marker(doc: DocumentT, token: str, hint: str = "") -> Paragraph:
    """Paragraphe-marqueur, remplacé par le renderer. Discret (gris italique).

    Returns:
        Paragraph: Le paragraphe-marqueur créé.
    """
    p = doc.add_paragraph()
    run = p.add_run(token + (f"  {hint}" if hint else ""))
    _set_run(run, size=9, italic=True, color="BFBFBF")
    return p


def _shade_cell(cell: _Cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _cell_text(
    cell: _Cell,
    text: str,
    *,
    bold: bool = False,
    align: WD_ALIGN_PARAGRAPH | None = None,
    color: str | None = None,
) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    _set_run(p.add_run(text), bold=bold, size=10, color=color)


def build(
    output_path: str | Path = config.DEFAULT_TEMPLATE,
) -> Path:
    """Génère un gabarit ``.docx`` (legacy v2) avec les marqueurs ``[[...]]``.

    Returns:
        Path: Le chemin du gabarit généré.
    """
    doc = Document()

    # Marges un peu resserrées.
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Pt(36)
        section.left_margin = section.right_margin = Pt(45)

    # 1. Titre
    _title_line(doc, "Centre des Opérations d'Urgence de Santé Publique", size=14)
    _title_line(doc, "« COUSP-RDC »", size=13)
    _title_line(doc, "Système de Gestion de l'Incident MVE-17", size=12)
    _title_line(
        doc,
        "Rapport de Situation de la 17ᵉ Épidémie de la Maladie à Virus Ebola / RDC",
        size=12,
    )
    _title_line(doc, "[[TITRE_NUMERO]]", size=12)
    doc.add_paragraph()

    # 2. Tableau d'identité
    t_id = doc.add_table(rows=len(IDENTITY_ROWS), cols=2)
    t_id.style = "Table Grid"
    for i, label in enumerate(IDENTITY_ROWS):
        _cell_text(t_id.cell(i, 0), label, bold=True)
        _shade_cell(t_id.cell(i, 0), "DEEAF6")
        _cell_text(t_id.cell(i, 1), "")  # rempli par le renderer
    doc.add_paragraph()

    # 3. Tableau des chiffres-clés (situation globale)
    _title_line(doc, "Situation globale", size=12, color="1F4E79")
    t_kpi = doc.add_table(rows=2, cols=len(KPI_LABELS))
    t_kpi.style = "Table Grid"
    for j, label in enumerate(KPI_LABELS):
        # Ligne 0 : valeur (placeholder), ligne 1 : libellé.
        _cell_text(
            t_kpi.cell(0, j),
            "—",
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            color=_RED,
        )
        _cell_text(t_kpi.cell(1, j), label, align=WD_ALIGN_PARAGRAPH.CENTER)
        _shade_cell(t_kpi.cell(1, j), "F2F2F2")
    doc.add_paragraph()

    # 4. Synthèse par province (en-têtes reproduits) ----------------
    _title_line(doc, "Répartition par province", size=12, color="1F4E79")
    t_prov = doc.add_table(rows=1, cols=len(PROVINCE_HEADERS))
    t_prov.style = "Table Grid"
    for j, h in enumerate(PROVINCE_HEADERS):
        _cell_text(t_prov.cell(0, j), h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _shade_cell(t_prov.cell(0, j), "DEEAF6")
    _marker(doc, "[[PROVINCE_ROWS]]", "(lignes par province ajoutées ici)")
    doc.add_paragraph()

    # Numérotation alignée sur le SitRep officiel (PDF) : Faits saillants en 1,
    # Contexte en 2, Analyse épidémiologique en 3.
    # 1. Faits saillants
    _heading(doc, "1. FAITS SAILLANTS")
    _marker(doc, "[[FAITS_SAILLANTS]]")

    # 2. Contexte
    _heading(doc, "2. CONTEXTE ÉPIDÉMIOLOGIQUE ET OPÉRATIONNEL")
    _marker(doc, "[[CONTEXTE]]")

    # 3. Analyse épidémiologique détaillée
    _heading(doc, "3. ANALYSE ÉPIDÉMIOLOGIQUE DÉTAILLÉE")
    _marker(doc, "[[TABLEAU_I]]", "(distribution spatiale + répartition par ZS)")
    _marker(doc, "[[CARTE]]", "(cartes province + zone de santé)")
    _marker(doc, "[[COURBE_EPI]]", "(courbe épidémique)")
    _marker(doc, "[[PYRAMIDE]]", "(pyramide âge/sexe)")
    _marker(doc, "[[TABLEAU_CROISE]]", "(répartition sexe x tranche d'âge)")
    _marker(doc, "[[TABLEAU_II]]", "(alertes de la période par zone de santé)")

    # 4. Actions de réponse
    _heading(doc, "4. ACTIONS DE RÉPONSE")
    for i, (title, token) in enumerate(ACTION_SUBSECTIONS, start=1):
        _subheading(doc, f"4.{i}. {title}")
        _marker(doc, token)

    # 5. Défis
    _heading(doc, "5. DÉFIS")
    _marker(doc, "[[DEFIS]]")

    # 6. Recommandations
    _heading(doc, "6. RECOMMANDATIONS")
    _marker(doc, "[[RECOMMANDATIONS]]")

    # 11. Contacts
    _heading(doc, "POUR TOUTE INFORMATION SUPPLÉMENTAIRE, VEUILLEZ CONTACTER")
    _marker(doc, "[[CONTACTS]]")

    output_path = Path(output_path)
    doc.save(str(output_path))
    return output_path


if __name__ == "__main__":
    path = build()
    print(f"Template propre généré : {path}")
