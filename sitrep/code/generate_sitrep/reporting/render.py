from __future__ import annotations

import re
from collections.abc import Callable
from pathlib import Path

import config
from data.model import SitRepData
from docx import Document
from docx.document import Document as DocumentT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml.etree import _Element
from reporting.highlights import build_highlights
from utils import docx as dx
from utils.dates import fr_date as _fr_date

Filler = Callable[[dx.Cursor], None]

_FONT = dx.FONT
_CENTER = dx.CENTER
_HEADER_FILL = dx.HEADER_FILL
_set_cell = dx.set_cell
_set_cell_shading = dx.set_cell_shading
_set_cell_border_bottom = dx.set_cell_border_bottom
_Cursor = dx.Cursor
_para = dx.para
_bullet = dx.bullet
_table = dx.table
_marker_paragraph = dx.marker_paragraph
_replace_marker = dx.replace_marker
_find_table = dx.find_table
_norm = dx.norm


def _image_filler(
    doc: DocumentT,
    path: str | Path | None,
    *,
    width_in: float = 6.0,
    caption: str | None = None,
) -> Filler:
    """Insère une image centrée (avec un libellé optionnel) sous le marqueur.

    Si le fichier est absent, écrit « (visuel indisponible) » à la place pour ne
    pas casser le rendu.

    Returns:
        Filler: La fonction de remplissage à passer à ``_replace_marker``.
    """

    def fill(cur: dx.Cursor) -> None:
        if caption:
            p = _para(doc, caption, bold=True, size=10, align=_CENTER, color="1F4E79")
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            cur.add(p._p)
        if not path or not Path(path).exists():
            cur.add(
                _para(
                    doc,
                    "(visuel indisponible)",
                    italic=True,
                    size=9,
                    color="BFBFBF",
                    align=_CENTER,
                )._p
            )
            return
        p = _para(doc, "", align=_CENTER)
        p.add_run().add_picture(str(path), width=Inches(width_in))
        cur.add(p._p)

    return fill


def _set_document_language(doc: DocumentT, lang: str = "fr-FR") -> None:
    """Définit le français comme langue d'édition/correction par défaut.

    Word lit la langue depuis ``w:docDefaults/w:rPrDefault/w:rPr/w:lang`` de
    ``styles.xml``. Le template hérite du défaut python-docx (``en-US``), d'où
    la détection « anglais ». On force ``lang`` (et la valeur East-Asia) pour
    tout le document, en créant les nœuds manquants au besoin.
    """
    styles = doc.styles.element
    node = styles.find(qn("w:docDefaults"))
    if node is None:
        node = OxmlElement("w:docDefaults")
        styles.insert(0, node)
    for tag in ("w:rPrDefault", "w:rPr", "w:lang"):
        child = node.find(qn(tag))
        if child is None:
            child = OxmlElement(tag)
            node.append(child)
        node = child
    node.set(qn("w:val"), lang)
    node.set(qn("w:eastAsia"), lang)


# Remplissages
def _fill_title_number(doc: DocumentT, data: SitRepData) -> None:
    """Renseigne le titre « SitRep N°{num}/{INCIDENT}_{date de rapportage} ».

    Le marqueur ``[[TITRE_NUMERO]]`` est désormais dans le cadran du titre (zone
    de texte d'une forme), présent en double via ``mc:AlternateContent``
    (``Choice`` DrawingML + ``Fallback`` VML) : on remplace le texte de chaque
    run portant le marqueur **sans recréer le run**, pour préserver la mise en
    forme appliquée à la main dans le gabarit. La date est la date de rapportage
    au format ``jj/mm/aaaa``.
    """
    titre = f"SitRep N°{data.sitrep_number}/{config.INCIDENT}_{data.reporting_end:%d/%m/%Y}"
    found = False
    for t in doc.element.iter(_TAG_T):
        if t.text and "[[TITRE_NUMERO]]" in t.text:
            t.text = t.text.replace("[[TITRE_NUMERO]]", titre)
            found = True
    if found:
        return
    # Repli : ancien gabarit où le marqueur était un paragraphe de corps.
    marker = _marker_paragraph(doc, "[[TITRE_NUMERO]]")
    if marker is None:
        return
    marker.text = ""
    run = marker.add_run(titre)
    run.bold = True
    run.font.name = _FONT
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(config.ACCENT_RED)


def _fill_identity(doc: DocumentT, data: SitRepData) -> None:
    t = _find_table(
        doc,
        lambda t: len(t.columns) == 2 and "provinces touchées" in t.cell(0, 0).text.strip().lower(),
    )
    if t is None:
        return
    za = data.zones_atteintes
    n_provinces = len(data.provinces_touchees)
    n_zones = sum(int(d.get("touchees", 0)) for d in za.values())

    def _zones_line(p: str) -> str:
        d = za.get(p, {})
        touchees, total = d.get("touchees", 0), d.get("total")
        compte = f"{touchees}/{total}" if total else str(touchees)
        zones = ", ".join(data.zones_by_province.get(p, []))
        suffixe = f" : {zones}" if zones else ""
        return f"• {p} ({compte}){suffixe}"

    zones_txt = "\n".join(_zones_line(p) for p in data.provinces_touchees)
    values: dict[str, tuple[str | None, str]] = {
        "provinces touchées": (
            f"Provinces touchées ({n_provinces})",
            ", ".join(data.provinces_touchees),
        ),
        "zones de santé touchées": (
            f"Zones de Santé touchées ({n_zones})",
            zones_txt,
        ),
        "date de rapportage": (None, data.reporting_label),
        "date de publication": (None, _fr_date(data.publication_date)),
    }
    for row in t.rows:
        label = row.cells[0].text.strip().lower()
        if label not in values:
            continue
        new_label, value = values[label]
        # Tout le tableau récap en gras, police 12 (fonds d'origine conservés).
        if new_label is not None:
            _set_cell(row.cells[0], new_label, bold=True, size=12)
        _set_cell(row.cells[1], value, bold=True, size=12)


def _fill_kpi(doc: DocumentT, data: SitRepData) -> None:
    kpi = data.kpi
    label_to_value = {
        "cumul cas confirmés": kpi["cumul_confirmes"],
        "cumul décès parmi les confirmés": kpi["cumul_deces"],
        "cas suspects en cours d'investigation": kpi["cumul_suspects"],
        "cas suspects en isolement": kpi["suspects_isolement"],
        "cas confirmés actifs": kpi["confirmes_actifs"],
        "guéris": kpi["gueris"],
        "taux de suivi de contacts": kpi["taux_suivi_contacts"],
    }
    t = _find_table(
        doc,
        lambda t: len(t.rows) == 2 and _norm(t.cell(1, 0).text) == "cumul cas confirmés",
    )
    if t is None:
        return
    for j in range(len(t.columns)):
        label = _norm(t.cell(1, j).text)
        if label in label_to_value:
            _set_cell(
                t.cell(0, j),
                label_to_value[label],  # type: ignore
                bold=True,
                align=_CENTER,
                color=config.ACCENT_RED,
            )


# --- Bandeau d'accueil : cartes KPI en formes/zones de texte (drawing) -------
# Les cartes de la 1re page ne sont pas des tableaux mais des formes (wps:wsp)
# nommées ; on injecte les valeurs calculées dans leurs zones de texte par nom,
# en préservant les runs annexes (astérisques en exposant, libellés « décès »…).
# Tags qualifiés (notation Clark) : python-docx surcharge ``.xpath`` (nsmap figé
# sans ``wps``) ; on passe donc par ``.iter()`` natif lxml.
_NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_NS_WPS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
_TAG_WSP = f"{{{_NS_WPS}}}wsp"
_TAG_CNVPR = f"{{{_NS_WPS}}}cNvPr"
_TAG_TXBX = f"{{{_NS_W}}}txbxContent"
_TAG_T = f"{{{_NS_W}}}t"


def _fr_int(value: object) -> str:
    """Formate un entier à la française (séparateur de milliers = espace insécable).

    Returns:
        str: La valeur formatée, ou la chaîne telle quelle si ce n'est pas un int
        (ex. ``config.ND``).
    """
    return f"{value:,}".replace(",", "\u00a0") if isinstance(value, int) else str(value)


def _fr_pct(value: float | str, decimals: int = 1) -> str:
    """Formate un pourcentage à la française (virgule décimale).

    Returns:
        str: ``"25,3%"`` ou la valeur telle quelle si non numérique (``ND``).
    """
    if not isinstance(value, (int, float)):
        return str(value)
    return f"{value:.{decimals}f}".replace(".", ",") + "%"


def _shape_by_name(doc: DocumentT, name: str) -> _Element | None:
    """Retourne la 1re forme (wps:wsp) portant ce nom, ou ``None``.

    Returns:
        L'élément XML de la forme, ou ``None`` si absente.
    """
    for sp in doc.element.iter(_TAG_WSP):
        cnv = next(sp.iter(_TAG_CNVPR), None)
        if cnv is not None and cnv.get("name") == name:
            return sp
    return None


def _shape_texts(shape: _Element) -> list[_Element]:
    """Retourne les nœuds ``w:t`` de la zone de texte de la forme, dans l'ordre.

    Returns:
        list: Les éléments ``w:t`` de tous les ``w:txbxContent`` de la forme.
    """
    ts: list[_Element] = []
    for txbx in shape.iter(_TAG_TXBX):
        ts.extend(txbx.iter(_TAG_T))
    return ts


def _set_first_run(doc: DocumentT, name: str, text: str) -> None:
    """Remplace le 1er run de la forme, laissant les suivants (ex. ``*`` exposant)."""
    sp = _shape_by_name(doc, name)
    if sp is None:
        return
    ts = _shape_texts(sp)
    if ts:
        ts[0].text = text


def _set_collapse(doc: DocumentT, name: str, text: str) -> None:
    """Met ``text`` dans le 1er run et vide les runs suivants (ex. « 67,0 » + « % »)."""
    sp = _shape_by_name(doc, name)
    if sp is None:
        return
    ts = _shape_texts(sp)
    if not ts:
        return
    ts[0].text = text
    for t in ts[1:]:
        t.text = ""


def _set_suspects(doc: DocumentT, name: str, n_text: str, deces_text: str) -> None:
    """Carte « Cas suspects du jour » : nombre (1er run) + décès (run avant « décès »).

    Conserve les runs « ( », « décès », « ) » et l'appel de note « ** ».
    """
    sp = _shape_by_name(doc, name)
    if sp is None:
        return
    ts = _shape_texts(sp)
    if not ts:
        return
    ts[0].text = n_text
    for i, t in enumerate(ts):
        if t.text and "décès" in t.text and i >= 1:
            ts[i - 1].text = deces_text  # le run du nombre de décès (« 30 »)
            break


def _update_provinces_label(doc: DocumentT, n: int) -> None:
    """Rend dynamique le « les N provinces » des libellés (accord singulier/pluriel)."""
    repl = "la province" if n == 1 else f"les {n} provinces"
    pat = re.compile(r"les\s+\d+\s+provinces?")
    for txbx in doc.element.iter(_TAG_TXBX):
        for t in txbx.iter(_TAG_T):
            if t.text and pat.search(t.text):
                t.text = pat.sub(repl, t.text)


def _fill_home_kpi(doc: DocumentT, data: SitRepData) -> None:
    """Injecte les indicateurs calculés dans les cartes KPI de la page d'accueil."""
    kpi = data.kpi
    conf = kpi["cumul_confirmes"]
    dec = kpi["cumul_deces"]
    cfr = round(int(dec) / int(conf) * 100, 1) if isinstance(conf, int) and conf else 0.0  # type: ignore

    # Cumul confirmés : 1er run = nombre, run suivant = « * » en exposant (à garder).
    _set_first_run(doc, "Valeur cumul cas confimrés", _fr_int(conf))
    # Décès + létalité : la zone compte plusieurs runs (nombre + « (xx,x%) ») sans
    # appel de note → on regroupe pour ne pas laisser l'ancien pourcentage.
    _set_collapse(
        doc,
        "Valeur cumul décès parmi les confirmés",
        f"{_fr_int(dec)}({_fr_pct(cfr)})",
    )
    # Patients en isolement/hospitalisation : effectif = cas actifs (stock) ;
    # le taux d'occupation global reste ND (à définir ultérieurement).
    _set_first_run(
        doc,
        "Valeur patients en isolement",
        f"{_fr_int(kpi['confirmes_isolement'])}({config.ND})",
    )
    _set_first_run(doc, "Valeur guéris", _fr_int(kpi["gueris"]))

    # Taux de suivi des contacts : « 67,0 » + « % » → « ND » (run « % » vidé).
    _set_collapse(doc, "Valeur taux de suivi des contacts", str(kpi["taux_suivi_contacts"]))

    # Cas suspects du jour : nombre + (décès parmi les suspects, ND pour l'instant).
    deces_jour = kpi["deces_suspects_jour"]
    deces_text = f"{_fr_int(deces_jour)} "  # espace final calqué sur le gabarit (« 30 »)
    _set_suspects(doc, "Valeur cas suspects décès", _fr_int(kpi["suspects_jour"]), deces_text)

    # Libellés « N provinces » dynamiques.
    _update_provinces_label(doc, len(data.provinces_touchees))


def _fill_province_rows(doc: DocumentT, data: SitRepData) -> None:
    t = _find_table(doc, lambda t: "guéris du jour" in t.cell(0, -1).text.strip().lower())
    if t is None:
        return
    for r in data.province_summary:
        cells = t.add_row().cells
        bold = r["province"] == "Total"
        # Colonnes : Province | Cas confirmés | Décès confirmés | Suspects en
        # cours d'investigation | Suspects en isolement | Cas confirmés actifs |
        # Guéris du jour.
        vals = [
            r["province"],
            r["confirmes"],
            r["deces"],
            r["suspects"],
            r["suspects_isolement"],
            r["actifs"],
            r["gueris"],
        ]
        for j, val in enumerate(vals[: len(cells)]):
            _set_cell(cells[j], val, bold=bold, align=_CENTER if j else None)
            if bold:
                _set_cell_shading(cells[j], _HEADER_FILL)

    # Filet de clôture explicite sur la dernière ligne (cf. utils.docx.table).
    if t.rows:
        for cell in t.rows[-1].cells:
            _set_cell_border_bottom(cell)

    # Le marqueur [[PROVINCE_ROWS]] (sous le tableau) n'a plus d'utilité.
    m = _marker_paragraph(doc, "[[PROVINCE_ROWS]]")
    if m is not None:
        m._p.getparent().remove(m._p)


def _insert_analytics(doc: DocumentT, data: SitRepData, charts: dict[str, Path | None]) -> None:
    def caption(cur: dx.Cursor, text: str) -> None:
        # Espace avant pour décoller le titre du tableau/visuel précédent.
        p = _para(doc, text, bold=True, size=10, align=_CENTER, color="1F4E79")
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        cur.add(p._p)

    def image(cur: dx.Cursor, path: str | Path | None) -> None:
        _image_filler(doc, path)(cur)

    rep = data.reporting_label
    rep_end = _fr_date(data.reporting_end)

    def _zones_cell(d: dict) -> str:
        k, n = d["zones_touchees"], d["zones_total"]
        if not n:
            return str(k)
        pct = f"{k / n * 100:.1f}".replace(".", ",")
        return f"{k} sur {n} ({pct}%)"

    def _pct(k: int, total: int) -> str:
        return f"{k / total * 100:.1f}".replace(".", ",") + "%" if total else config.ND

    def fill_tableau1(cur: dx.Cursor) -> None:
        # 3.1 Distribution spatiale (par province) — style PDF, avec CFR.
        caption(cur, f"Tableau I : Distribution spatiale au {rep_end}")
        rows = [
            [
                d["province"],
                d["confirmes"],
                d["deces"],
                f"{d['cfr']:.1f}".replace(".", ",") + "%",
                _zones_cell(d),
                d["nouveaux"],
            ]
            for d in data.distribution_spatiale
        ]
        cur.add(
            _table(
                doc,
                [
                    "Province",
                    "Cas confirmés",
                    "Décès confirmés",
                    "CFR",
                    "Zones de santé touchées",
                    f"Nouveaux cas ({rep})",
                ],
                rows,
            )._tbl
        )

    def fill_tableau_zone(cur: dx.Cursor) -> None:
        # 3.2 Répartition détaillée par zone de santé.
        caption(cur, f"Tableau II : Répartition par zone de santé au {rep_end}")
        zrows = [
            [
                r["province"],
                r["zone"],
                r["confirmes"],
                r["deces"],
                r["actifs"],
                r["gueris"],
            ]
            for r in data.tableau1
        ]
        tot = data.tableau1_total
        zrows.append(["Total", "", tot["confirmes"], tot["deces"], tot["actifs"], tot["gueris"]])
        cur.add(
            _table(
                doc,
                [
                    "Province",
                    "Zone de santé",
                    "Cas confirmés",
                    "Décès confirmés",
                    "Cas confirmés actifs",
                    "Guéris",
                ],
                zrows,
            )._tbl
        )

    def fill_carte(cur: dx.Cursor) -> None:
        caption(cur, "Cas confirmés MVE par province")
        image(cur, charts.get("province_situation_map"))
        caption(cur, "Spatialisation des cas confirmés MVE par zone de santé")
        image(cur, charts.get("zone_situation_map"))

    def fill_courbe(cur: dx.Cursor) -> None:
        caption(
            cur,
            "Cas confirmés par semaine épidémiologique de début des symptômes",
        )
        image(cur, charts.get("epi_curve"))

    def fill_courbe_symptome(cur: dx.Cursor) -> None:
        caption(
            cur,
            "Combinaisons de symptômes — Cas confirmés",
        )
        image(cur, charts.get("combi_symptomes"))

    def fill_pyramide(cur: dx.Cursor) -> None:
        caption(cur, "Cas confirmés par tranche d'âge et sexe")
        image(cur, charts.get("age_sex_pyramid"))

    def fill_croise(cur: dx.Cursor) -> None:
        caption(cur, "Répartition des cas confirmés par sexe et tranche d'âge")
        rows = [
            [r["tranche_age"], r["Masculin"], r["Feminin"], r["Total"]]
            for r in data.agesex_crosstab
        ]
        cur.add(_table(doc, ["Tranche d'âge", "Masculin", "Feminin", "Total"], rows)._tbl)

        # Note de complétude : les cas confirmés sans sexe ou tranche d'âge
        # renseigné sont exclus de ce tableau, d'où un total < cas confirmés de
        # la situation globale. Basée sur l'écart réel, elle ne s'affiche que
        # lorsqu'il existe. À retirer une fois l'alignement avec les indicateurs
        # produits par le programme effectué.
        croise_total = next(
            (r["Total"] for r in data.agesex_crosstab if r["tranche_age"] == "Total"),
            sum(r["Total"] for r in data.agesex_crosstab),
        )
        cumul_confirmes = data.kpi.get("cumul_confirmes")
        n_exclus = cumul_confirmes - croise_total if isinstance(cumul_confirmes, int) else 0
        if n_exclus > 0:
            accord = "cas confirmé" if n_exclus == 1 else "cas confirmés"
            note = (
                f"Note : {n_exclus} {accord} sans sexe ou tranche d'âge renseigné sont "
                "exclus de cette répartition ; le total par tranche d'âge et sexe diffère "
                "donc du nombre de cas confirmés de la situation globale, en attendant "
                "l'alignement avec les indicateurs produits par le programme."
            )
            p = _para(doc, note, italic=True, size=8, color="808080")
            p.paragraph_format.space_before = Pt(2)
            cur.add(p._p)

    def fill_surveillance(cur: dx.Cursor) -> None:
        # 4.2 — Tableau III : gestion des alertes (période).
        caption(cur, f"Tableau III : Indicateurs de surveillance ({rep})")
        s = data.surveillance_indics
        tot = s["alertes_remontees"]
        rows = [
            ["Alertes remontées", s["alertes_remontees"], config.ND, config.ND],
            [
                "Alertes investiguées",
                f"{s['alertes_investiguees']} ({_pct(s['alertes_investiguees'], tot)})",
                "≥ 95%",
                config.ND,
            ],
            [
                "Alertes validées",
                f"{s['alertes_validees']} ({_pct(s['alertes_validees'], tot)})",
                config.ND,
                config.ND,
            ],
            ["Nouvelles zones de santé touchées", s["nouvelles_zones"], "0", config.ND],
        ]
        cur.add(_table(doc, ["Indicateur", "Valeur", "Cible", "Performance"], rows)._tbl)

    def fill_contacts(cur: dx.Cursor) -> None:
        # 4.2 — Tableau IV : suivi des contacts (non disponible -> à saisir).
        caption(cur, f"Tableau IV : Suivi des contacts des cas confirmés au {rep_end}")
        rows = [
            [lab, config.ND]
            for lab in (
                "Contacts listés",
                "Contacts sous suivi",
                "Contacts vus du jour",
                "Taux de suivi des contacts",
                "Contacts sortis de suivi",
            )
        ]
        cur.add(_table(doc, ["Indicateur", "Valeur"], rows)._tbl)

    def fill_labo(cur: dx.Cursor) -> None:
        # 4.3 — Tableau V : indicateurs laboratoire (période).
        caption(cur, f"Tableau V : Indicateurs laboratoire ({rep})")
        lab = data.labo_indics
        rows = [
            ["Échantillons collectés", lab["collectes"], config.ND],
            [
                "Échantillons analysés",
                lab["analyses"],
                f"{_pct(lab['analyses'], lab['collectes'])} des collectés",
            ],
            [
                "Échantillons positifs",
                lab["positifs"],
                f"Taux de positivité {_pct(lab['positifs'], lab['analyses'])}",
            ],
            ["Échantillons en cours d'analyse", lab["en_cours"], config.ND],
            [
                "Échantillons positifs (cas suspects décédés)",
                lab["positifs_suspect_decedes"],
                config.ND,
            ],
        ]
        cur.add(_table(doc, ["Indicateur", "Valeur", "Analyse"], rows)._tbl)

    def fill_prise_en_charge(cur: dx.Cursor) -> None:
        # 4.5 — Tableau VI : indicateurs de prise en charge.
        caption(cur, f"Tableau VI : Indicateurs de prise en charge au {rep_end}")
        pc = data.prise_en_charge_indics
        rows = [
            ["Cas confirmés en isolement", pc["confirmes_isolement"]],
            ["Cas suspects en isolement", pc["suspects_isolement"]],
            ["Cas confirmés actifs", pc["actifs"]],
            ["Guéris du jour", pc["gueris_jour"]],
            ["Létalité hospitalière", pc["letalite"]],
            ["Lits CTE disponibles vs occupés", pc["lits"]],
        ]
        cur.add(_table(doc, ["Indicateur", "Valeur"], rows)._tbl)

    def fill_mouvement(cur: dx.Cursor) -> None:
        # 4.5 — Tableau VII : mouvement des patients (non disponible -> à saisir).
        caption(
            cur,
            f"Tableau VII : Mouvement des malades dans les établissements de soins au {rep_end}",
        )
        rows = [
            [lab, config.ND]
            for lab in (
                "Malades au lit (report veille)",
                "Nouvelles admissions du jour",
                "Sorties (guéris / décédés / transférés)",
                "Malades en isolement",
            )
        ]
        cur.add(_table(doc, ["Mouvement", "Valeur"], rows)._tbl)

    _replace_marker(doc, "[[TABLEAU_I]]", fill_tableau1)
    _replace_marker(doc, "[[TABLEAU_II]]", fill_tableau_zone)
    _replace_marker(doc, "[[CARTE]]", fill_carte)
    _replace_marker(doc, "[[COURBE_EPI]]", fill_courbe)
    _replace_marker(doc, "[[COURBE_EPI_SYMPTOME]]", fill_courbe_symptome)
    _replace_marker(doc, "[[PYRAMIDE]]", fill_pyramide)
    _replace_marker(doc, "[[TABLEAU_CROISE]]", fill_croise)
    _replace_marker(doc, "[[TABLEAU_III]]", fill_surveillance)
    _replace_marker(doc, "[[TABLEAU_IV]]", fill_contacts)
    _replace_marker(doc, "[[TABLEAU_V]]", fill_labo)
    _replace_marker(doc, "[[TABLEAU_VI]]", fill_prise_en_charge)
    _replace_marker(doc, "[[TABLEAU_VII]]", fill_mouvement)


def _bullets_filler(doc: DocumentT, items: list[str] | None) -> Filler:
    def fill(cur: dx.Cursor) -> None:
        if not items:
            cur.add(_para(doc, "À compléter.", italic=True, size=10, color="808080")._p)
            return
        for it in items:
            cur.add(_bullet(doc, str(it))._p)

    return fill


def _paragraphs_filler(doc: DocumentT, value: str | list[str] | None) -> Filler:
    """Rend un texte continu (str ou liste de str) en paragraphes justifiés.

    Returns:
        Filler: La fonction de remplissage à passer à ``_replace_marker``.
    """
    if isinstance(value, str):
        items = [value]
    else:
        items = list(value or [])

    def fill(cur: dx.Cursor) -> None:
        if not items:
            cur.add(_para(doc, "À compléter.", italic=True, size=10, color="808080")._p)
            return
        for it in items:
            p = _para(doc, str(it), size=10)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            cur.add(p._p)

    return fill


def _highlights_filler(doc: DocumentT, auto: list[str], items: list[str] | None) -> Filler:
    """Faits saillants auto (« à date ») suivis du narratif manuel.

    Returns:
        Filler: La fonction de remplissage à passer à ``_replace_marker``.
    """

    def fill(cur: dx.Cursor) -> None:
        produced = False
        for line in auto:
            cur.add(_bullet(doc, line)._p)
            produced = True
        for it in items or []:
            cur.add(_bullet(doc, str(it))._p)
            produced = True
        if not produced:
            cur.add(_para(doc, "À compléter.", italic=True, size=10, color="808080")._p)

    return fill


def _inject_narrative(doc: DocumentT, narrative: dict, data: SitRepData) -> None:
    narrative = narrative or {}
    actions = narrative.get("actions_reponse", {}) or {}

    _replace_marker(doc, "[[CONTEXTE]]", _paragraphs_filler(doc, narrative.get("contexte")))
    # [[FIGURE_1]] : chronologie des faits de l'épidémie, fournie en image.
    _replace_marker(
        doc,
        "[[FIGURE_1]]",
        _image_filler(
            doc,
            config.IMAGE_CHRONOLOGIE,
            caption="Chronologie des faits de l'épidémie de MVE (Ebola-Bundibugyo)",
        ),
    )
    _replace_marker(
        doc,
        "[[FAITS_SAILLANTS]]",
        _highlights_filler(doc, build_highlights(data), narrative.get("faits_saillants")),
    )
    for token, key in [
        ("[[ACTIONS_COORDINATION]]", "coordination"),
        ("[[ACTIONS_PRISE_EN_CHARGE]]", "prise_en_charge"),
        ("[[ACTIONS_CREC]]", "crec"),
        ("[[ACTIONS_PCI_WASH]]", "pci_wash"),
        ("[[ACTIONS_LOGISTIQUE]]", "logistique"),
        ("[[ACTIONS_SECURITE]]", "securite"),
    ]:
        _replace_marker(doc, token, _bullets_filler(doc, actions.get(key)))
    _replace_marker(doc, "[[DEFIS]]", _bullets_filler(doc, narrative.get("defis")))
    _replace_marker(
        doc,
        "[[RECOMMANDATIONS]]",
        _bullets_filler(doc, narrative.get("recommandations")),
    )
    # [[CONTACTS]] : bloc « info-contacts » fourni en image par le programme.
    _replace_marker(doc, "[[CONTACTS]]", _image_filler(doc, config.IMAGE_CONTACTS))


def render(
    data: SitRepData,
    charts: dict[str, Path | None],
    template_path: str | Path,
    output_path: str | Path,
    narrative: dict | None = None,
) -> Path:
    """Produit le fichier SitRep .docx et renvoie son chemin.

    Returns:
        Path: Le chemin du ``.docx`` généré.
    """
    doc = Document(str(template_path))

    _set_document_language(doc)
    _fill_title_number(doc, data)
    _fill_identity(doc, data)
    _fill_kpi(doc, data)
    _fill_home_kpi(doc, data)
    _fill_province_rows(doc, data)
    _insert_analytics(doc, data, charts)
    _inject_narrative(doc, narrative or {}, data)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
